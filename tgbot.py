import os
from telegram import Update, InputFile
from telegram.ext import Application, CommandHandler, MessageHandler, CallbackContext, ConversationHandler, filters
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# Conversation states
TITLE, HEADING, TEXT = range(3)

# Step 1: Define a function to create a Word document with a specific format
def create_word_document(file_name: str, heading: str, text: str):
    """
    Create a Word document with predefined formatting.

    Args:
        file_name (str): The name of the output Word document.
        heading (str): The heading of the document.
        text (str): The text to include in the document.
    """
    doc = Document()

    # Adding a heading to the document
    title = doc.add_heading(heading, level=1)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Adding the text with custom styling
    for line in text.split('\n'):
        if line.startswith("###"):
            doc.add_heading(line.strip("# "), level=3)
        elif line.startswith("---"):
            paragraph = doc.add_paragraph()
            run = paragraph.add_run(line.strip("- "))
            run.bold = True
        else:
            paragraph = doc.add_paragraph(line)
            if paragraph.runs:  # Check if runs exist before accessing
                run = paragraph.runs[0]
                run.font.name = 'Arial'
                run.font.size = Pt(12)

    # Save the document
    doc.save(file_name)


# Step 2: Conversation Handlers
async def start(update: Update, context: CallbackContext):
    await update.message.reply_text("Welcome! What should the document be named?")
    return TITLE

async def set_title(update: Update, context: CallbackContext):
    context.user_data['file_name'] = f"{update.message.text}.docx"
    await update.message.reply_text("Great! Now, what should the heading of the document be?")
    return HEADING

async def set_heading(update: Update, context: CallbackContext):
    context.user_data['heading'] = update.message.text
    await update.message.reply_text("Perfect! Please send me the text for the document.")
    return TEXT

async def set_text(update: Update, context: CallbackContext):
    context.user_data['text'] = update.message.text
    file_name = context.user_data['file_name']
    heading = context.user_data['heading']
    text = context.user_data['text']

    # Create the Word document
    create_word_document(file_name, heading, text)

    # Send the document back to the user
    with open(file_name, 'rb') as document:
        await update.message.reply_document(InputFile(document, filename=file_name))

    # Clean up the temporary file
    if os.path.exists(file_name):
        os.remove(file_name)

    await update.message.reply_text("Your document has been created!")
    return ConversationHandler.END

async def cancel(update: Update, context: CallbackContext):
    await update.message.reply_text("Operation canceled.")
    return ConversationHandler.END

# Step 3: Define the main function to set up the bot
def main():
    TOKEN = '7299744318:AAHNfHh5ZMgAbdXv2guZaTSm4ZK-SHMp2KM'

    # Create an Application object
    application = Application.builder().token(TOKEN).build()

    # Create a conversation handler
    conv_handler = ConversationHandler(
        entry_points=[CommandHandler("start", start)],
        states={
            TITLE: [MessageHandler(filters.TEXT & ~filters.COMMAND, set_title)],
            HEADING: [MessageHandler(filters.TEXT & ~filters.COMMAND, set_heading)],
            TEXT: [MessageHandler(filters.TEXT & ~filters.COMMAND, set_text)],
        },
        fallbacks=[CommandHandler("cancel", cancel)],
    )

    # Add the conversation handler to the application
    application.add_handler(conv_handler)

    # Start the bot
    print("Bot is running...")
    application.run_polling()

# Step 4: Run the bot
if __name__ == "__main__":
    main()
